import os
import re
import logging
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document as DocxDocument  # To avoid clash with local 'Document'
import PyPDF2
import yaml  # PyYAML

# --- Configuration ---
INPUT_FOLDER = "input_assessments"
OUTPUT_FOLDER = "output_feedback"
MASTER_PROMPT_FILE = "master_prompt.txt"
LOG_FILE = "grading_process.log"
SUMMARY_FILE = "grading_summary.csv"
GRADE_REVIEW_PROMPT_FILE = "grade_review_prompt.txt"

# Setup basic logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(),  # Also print to console
    ],
)

# --- Helper Functions ---


def load_api_key():
    """Loads Gemini API key from .env file."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logging.error("GEMINI_API_KEY not found in .env file.")
        raise ValueError("API Key not configured.")
    return api_key


def get_student_name_from_filename(filename):
    """
    Attempts to extract a student name from filename.
    Example: "JohnDoe_Assignment1.docx" -> "JohnDoe"
    """
    name_part = os.path.splitext(filename)[0]
    # Simple heuristic: look for parts separated by common delimiters
    # that might indicate a name. This can be improved.
    potential_name = re.split(r"[_\-\s.]", name_part)[0]
    # Check if it looks like a name (e.g., starts with capital)
    if potential_name and potential_name[0].isupper():
        # Further refine if needed, e.g. look for CamelCase or multiple capitalized words
        return potential_name
    return None


def extract_text_from_docx(doc):
    """Extract text from paragraphs and tables in a DOCX Document."""
    text_parts = []
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        text_parts.append(para.text)
    return "\n".join(text_parts)


def extract_text_from_file(filepath):
    """Extracts text and author metadata from supported files."""
    _, extension = os.path.splitext(filepath)
    text = ""
    doc_author = None
    try:
        if extension.lower() == ".docx":
            doc = DocxDocument(filepath)
            doc_author = doc.core_properties.author or None
            text = extract_text_from_docx(doc)
        elif extension.lower() == ".pdf":
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    logging.warning(
                        f"PDF '{filepath}' is encrypted. Attempting to read anyway if default password allows."
                    )
                    # You might need to handle decryption if password is known: reader.decrypt('')
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
        else:  # Attempt plain text for other files
            logging.info(f"Attempting to read '{filepath}' as plain text.")
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        text = re.sub(r"\s{3,}", "\n\n", text).strip()
        if not text.strip():
            logging.warning(f"No text extracted or file is empty: {filepath}")
            return None, None
        return text, doc_author

    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
        return None, None
    except PyPDF2.errors.PdfReadError:
        logging.error(
            f"Could not read PDF (possibly corrupted or password protected): {filepath}"
        )
        return None, None
    except Exception as e:
        logging.error(f"Error extracting text from {filepath}: {e}")
        return None, None


def load_master_prompt():
    """Loads the master prompt template from file."""
    try:
        with open(MASTER_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Master prompt file '{MASTER_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading master prompt file: {e}")
        raise


def load_grade_review_prompt_template():
    """Loads the grade review prompt template from file."""
    try:
        with open(GRADE_REVIEW_PROMPT_FILE, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logging.error(f"Grade review prompt file '{GRADE_REVIEW_PROMPT_FILE}' not found.")
        raise
    except Exception as e:
        logging.error(f"Error reading grade review prompt file: {e}")
        raise


def construct_full_prompt(
    student_text, master_prompt_template, *, raise_on_missing=True
):
    """Insert student text into the master prompt template.

    Parameters
    ----------
    student_text : str
        The extracted text of the student's submission.
    master_prompt_template : str
        The prompt template loaded from ``master_prompt.txt``.
    raise_on_missing : bool, optional
        If ``True`` (default), a ``ValueError`` is raised when the required
        placeholder is not present in ``master_prompt_template``. The student
        text is still appended to the returned prompt to avoid an empty
        submission.
    """

    placeholder = "{{STUDENT_SUBMISSION_TEXT_HERE}}"

    if placeholder not in master_prompt_template:
        warning_msg = (
            f"Placeholder '{placeholder}' not found in master prompt template. "
            "Student text will be appended to the end of the prompt."
        )
        logging.warning(warning_msg)

        fallback_prompt = (
            master_prompt_template
            + "\n\n### STUDENT_SUBMISSION_TEXT_TO_GRADE:\n"
            + student_text
        )

        if raise_on_missing:
            # Raising an exception alerts the operator to fix the prompt file
            raise ValueError(warning_msg)

        return fallback_prompt

    return master_prompt_template.replace(placeholder, student_text)


def call_gemini_api(prompt_text, api_key):
    """Calls the Gemini API and returns the response text."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash-latest")  # Or your preferred model
    # Safety settings can be adjusted if needed
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]
    try:
        logging.info("Sending request to Gemini API...")
        # Add a timeout to generation_config if needed
        response = model.generate_content(prompt_text, safety_settings=safety_settings)
        # Check for empty or blocked responses
        if not response.parts:
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                logging.error(
                    f"Gemini API request blocked. Reason: {response.prompt_feedback.block_reason_message}"
                )
                return None
            else:
                logging.error("Gemini API returned an empty response with no parts.")
                return None

        # Assuming the response is text. If it could be multi-modal, access response.text
        ai_response_text = response.text
        logging.info("Received response from Gemini API.")
        return ai_response_text
    except Exception as e:
        logging.error(f"Gemini API call failed: {e}")
        return None


def parse_gemini_yaml_response(response_text):
    """Parses the YAML response from Gemini."""
    if not response_text:
        return None
    try:
        # LLMs can sometimes add markdown backticks around YAML
        cleaned_response = response_text.strip()
        if cleaned_response.startswith("```yaml"):
            cleaned_response = cleaned_response[7:]
        if cleaned_response.startswith("```"):
            cleaned_response = cleaned_response[3:]
        if cleaned_response.endswith("```"):
            cleaned_response = cleaned_response[:-3]

        cleaned_response = cleaned_response.strip()

        parsed_data = yaml.safe_load(cleaned_response)
        # Basic validation of expected structure
        if "assistant_reasons" in parsed_data and "assistant_grade" in parsed_data:
            return parsed_data
        else:
            logging.error(
                f"Parsed YAML does not have expected structure. Parsed: {parsed_data}"
            )
            return None
    except yaml.YAMLError as e:
        logging.error(
            f"Failed to parse YAML response from Gemini: {e}\nRaw response:\n{response_text}"
        )
        return None
    except Exception as e:
        logging.error(
            f"An unexpected error occurred during YAML parsing: {e}\nRaw response:\n{response_text}"
        )
        return None


def compute_overall_grade(breakdown):
    """Compute a letter grade from rubric breakdown points."""
    if not isinstance(breakdown, dict):
        return "N/A"

    total_possible = 25
    total_points = 0
    for item in breakdown.values():
        try:
            total_points += int(item.get("points", 0))
        except Exception:
            continue

    # Grade cutoffs are now based directly on raw points rather than a
    # simple ratio. 20–25 points = A, 15–19 = B, 12–14 = C. Values below 12
    # retain the previous mapping for D and E bands.
    if total_points >= 20:
        return "A"
    if total_points >= 15:
        return "B"
    if total_points >= 12:
        return "C"

    ratio = total_points / float(total_possible)
    band = round(ratio * 5)
    band = max(1, min(5, band))

    grading_scale = {"2": "D", "1": "E"}
    return grading_scale.get(str(band), "E")


def review_grade(student_text, grade_yaml_text, api_key, review_prompt_template=None):
    """Sends student text and the AI's grade to Gemini for fairness review."""
    if review_prompt_template is None:
        try:
            review_prompt_template = load_grade_review_prompt_template()
        except Exception:
            return None

    prompt = review_prompt_template
    if "{{STUDENT_SUBMISSION_TEXT_HERE}}" in prompt:
        prompt = prompt.replace("{{STUDENT_SUBMISSION_TEXT_HERE}}", student_text)
    else:
        logging.warning("Student submission placeholder missing in grade review prompt template")
        prompt += f"\n\nSTUDENT SUBMISSION:\n{student_text}"

    if "{{AI_GRADE_YAML_HERE}}" in prompt:
        prompt = prompt.replace("{{AI_GRADE_YAML_HERE}}", grade_yaml_text)
    else:
        logging.warning("AI grade placeholder missing in grade review prompt template")
        prompt += f"\n\nAI GRADE:\n{grade_yaml_text}"

    return call_gemini_api(prompt, api_key)


def extract_new_grade_from_review(review_text):
    """Attempt to extract a revised overall grade from the review text."""
    if not review_text:
        return None

    # Common patterns such as "grade should be B" or "recommended grade: A"
    patterns = [
        r"grade\s*should\s*be\s*([A-E])",
        r"recommended\s*grade[:\s]+([A-E])",
        r"proposed\s*grade[:\s]+([A-E])",
        r"new\s*grade[:\s]+([A-E])",
        r"should\s*be\s*an?\s*([A-E])",
    ]

    for pat in patterns:
        m = re.search(pat, review_text, re.IGNORECASE)
        if m:
            return m.group(1).upper()
    return None


def apply_criteria_adjustments(parsed_data, adjustments):
    """Apply band changes from review to the parsed YAML data."""
    if not adjustments:
        return

    grade_section = parsed_data.get("assistant_grade", {})
    breakdown = grade_section.get("breakdown", {})
    for crit, new_band in adjustments.items():
        if crit in breakdown:
            breakdown[crit]["band"] = new_band
            breakdown[crit]["points"] = new_band

    # Update total points if any changes were applied
    try:
        total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
        grade_section["total_points"] = total_points
    except Exception:
        pass

def extract_criteria_adjustments(review_text):
    """Parse review text for suggested band adjustments for specific criteria."""
    if not review_text:
        return {}

    criteria_aliases = {
        "symptom_analysis": [
            "symptom analysis",
            "knowledge & symptom analysis",
            "knowledge and symptom analysis",
            "criterion 1",
        ],
        "bps_factors": [
            "bps factors",
            "b-p-s factors",
            "biological, psychological & social factors",
            "criterion 2",
        ],
        "diagnostic_primary": [
            "diagnostic primary",
            "primary diagnosis accuracy",
            "primary diagnosis",
            "criterion 3",
        ],
        "diagnostic_diff": [
            "differential diagnosis reasoning",
            "differential diagnosis",
            "criterion 4",
        ],
        "treatment": ["treatment selection", "treatment", "criterion 5"],
        "communication": ["communication & referencing", "communication", "criterion 6"],
    }

    adjustments = {}
    for line in review_text.splitlines():
        lower_line = line.lower()
        for cid, aliases in criteria_aliases.items():
            if any(alias in lower_line for alias in aliases):
                numbers = re.findall(r"([1-5])", line)
                if numbers:
                    try:
                        new_band = int(numbers[-1])
                        adjustments[cid] = new_band
                    except ValueError:
                        pass
                break
    return adjustments


def format_feedback_as_docx(
    yaml_data, output_filepath, student_identifier, doc_author=None, override_grade=None
):
    """Formats the YAML data into a human-readable DOCX report."""
    try:
        doc = DocxDocument()
        doc.add_heading(f"Feedback Report for: {student_identifier}", level=1)
        if doc_author:
            doc.add_paragraph(f"Author (from file metadata): {doc_author}")

        # Overall Grade and Points
        grade_info = yaml_data.get("assistant_grade", {})
        breakdown = grade_info.get("breakdown", {})
        overall_grade = compute_overall_grade(breakdown)
        if override_grade:
            overall_grade = override_grade
        try:
            total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
        except Exception:
            total_points = grade_info.get("total_points", "N/A")
        max_total_points = 25

        doc.add_heading("Overall Assessment", level=2)
        doc.add_paragraph(f"Overall Grade: {overall_grade}")
        doc.add_paragraph(f"Total Points: {total_points} / {max_total_points}")
        doc.add_paragraph()  # Spacer

        # Criteria Breakdown
        doc.add_heading("Detailed Breakdown by Criterion", level=2)
        reasons = yaml_data.get("assistant_reasons", [])

        # You'll need to map criterion IDs from 'assistant_reasons' (e.g., 'symptom_analysis')
        # to their full names and max points if you want to display them nicely.
        # This mapping can be hardcoded or derived from your rubric JSON if it were loaded.
        # For simplicity, we'll use the IDs for now.
        # Example mapping (you'd get this from your RUBRIC_JSON ideally)
        rubric_criteria_details = {
            "symptom_analysis": {
                "name": "Knowledge & Symptom Analysis",
                "max_points": 5,
            },
            "bps_factors": {
                "name": "Biological, Psychological & Social Factors",
                "max_points": 4,
            },
            "diagnostic_primary": {
                "name": "Primary Diagnosis Accuracy & Justification",
                "max_points": 4,
            },
            "diagnostic_diff": {
                "name": "Differential Diagnosis Reasoning",
                "max_points": 4,
            },
            "treatment": {
                "name": "Treatment Selection & Justification",
                "max_points": 5,
            },
            "communication": {"name": "Communication & Referencing", "max_points": 3},
        }

        for idx, reason_item in enumerate(reasons, start=1):
            criterion_id = reason_item.get("criterion", "Unknown Criterion")
            band = reason_item.get("band", "N/A")
            rationale = reason_item.get("rationale", "No rationale provided.")
            evidence = reason_item.get("evidence", "No evidence quoted.")

            criterion_details = rubric_criteria_details.get(
                criterion_id,
                {"name": criterion_id.replace("_", " ").title(), "max_points": "N/A"},
            )
            criterion_name = criterion_details["name"]

            criterion_grade_info = breakdown.get(criterion_id, {})
            points_achieved = criterion_grade_info.get("points", "N/A")
            max_criterion_points = criterion_details["max_points"]

            doc.add_heading(f"{idx}. {criterion_name}", level=3)

            doc.add_paragraph(f"Band Achieved: {band}")
            doc.add_paragraph(f"Points: {points_achieved} / {max_criterion_points}")

            doc.add_paragraph("AI's Rationale:", style="Intense Quote")
            doc.add_paragraph(rationale)

            doc.add_paragraph("Evidence from Student's Work:", style="Intense Quote")
            if "\n" in evidence:
                for line in evidence.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(evidence if evidence else "N/A")

            doc.add_paragraph()  # Spacer

        doc.save(output_filepath)
        logging.info(f"Feedback report saved to: {output_filepath}")

    except Exception as e:
        logging.error(f"Failed to create DOCX report for {student_identifier}: {e}")


# --- Main Processing Logic ---
def main():
    logging.info("Starting AI Student Assessment Grader...")
    try:
        api_key = load_api_key()
        master_prompt_template = load_master_prompt()
    except Exception as e:
        logging.critical(f"Initialization failed: {e}")
        return

    if not os.path.exists(INPUT_FOLDER):
        logging.error(f"Input folder '{INPUT_FOLDER}' not found.")
        return
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        logging.info(f"Created output folder: {OUTPUT_FOLDER}")

    processed_files = 0
    successful_grades = 0
    summary_entries = []

    for filename in os.listdir(INPUT_FOLDER):
        filepath = os.path.join(INPUT_FOLDER, filename)
        if not os.path.isfile(filepath):
            continue  # Skip directories

        logging.info(f"--- Processing file: {filename} ---")
        processed_files += 1

        student_name_guess = get_student_name_from_filename(filename)
        student_identifier = (
            student_name_guess if student_name_guess else os.path.splitext(filename)[0]
        )

        extracted_text, doc_author = extract_text_from_file(filepath)
        if not extracted_text:
            logging.warning(
                f"Skipping {filename} due to text extraction failure or empty content."
            )
            continue

        # Simple word count for info, AI will use its own logic based on rubric
        word_count = len(extracted_text.split())
        logging.info(f"Extracted approx. {word_count} words from {filename}.")
        if word_count < 50:  # Arbitrary threshold for very short/empty files
            logging.warning(
                f"Extracted text for {filename} is very short ({word_count} words). May not be suitable for grading."
            )
            # continue # Optional: skip very short files

        full_prompt = construct_full_prompt(extracted_text, master_prompt_template)

        # For debugging, you might want to save the full prompt
        # with open(os.path.join(OUTPUT_FOLDER, f"{student_identifier}_prompt.txt"), "w", encoding="utf-8") as pf:
        #    pf.write(full_prompt)

        api_response = call_gemini_api(full_prompt, api_key)
        if not api_response:
            logging.warning(f"Skipping {filename} due to Gemini API call failure.")
            continue

        parsed_data = parse_gemini_yaml_response(api_response)
        if not parsed_data:
            logging.warning(f"Skipping {filename} due to YAML parsing failure.")
            # Save raw response for debugging
            raw_response_path = os.path.join(
                OUTPUT_FOLDER, f"{student_identifier}_raw_gemini_response.txt"
            )
            with open(raw_response_path, "w", encoding="utf-8") as f:
                f.write(api_response if api_response else "No response received.")
            logging.info(f"Raw Gemini response saved to: {raw_response_path}")
            continue

        output_filename_base = student_identifier
        output_docx_path = os.path.join(
            OUTPUT_FOLDER, f"{output_filename_base}_graded.docx"
        )

        review_text = review_grade(extracted_text, api_response, api_key)
        override_grade = None
        if review_text:
            override_grade = extract_new_grade_from_review(review_text)
            review_path = os.path.join(
                OUTPUT_FOLDER, f"{output_filename_base}_grade_review.txt"
            )
            try:
                with open(review_path, "w", encoding="utf-8") as rf:
                    rf.write(review_text)
                logging.info(f"Grade review saved to: {review_path}")
                if override_grade:
                    logging.info(
                        f"Applying grade override from review: {override_grade}"
                    )
                adjustments = extract_criteria_adjustments(review_text)
                if adjustments:
                    apply_criteria_adjustments(parsed_data, adjustments)
                    logging.info(f"Applied criterion adjustments: {adjustments}")
            except Exception as e:
                logging.error(f"Failed to save grade review for {student_identifier}: {e}")

        breakdown = parsed_data.get("assistant_grade", {}).get("breakdown", {})
        try:
            total_points = sum(int(item.get("points", 0)) for item in breakdown.values())
        except Exception:
            total_points = parsed_data.get("assistant_grade", {}).get("total_points", "N/A")
        overall_grade = compute_overall_grade(breakdown)
        if override_grade:
            overall_grade = override_grade

        summary_entries.append((student_identifier, total_points, overall_grade))

        format_feedback_as_docx(
            parsed_data,
            output_docx_path,
            student_identifier,
            doc_author=doc_author,
            override_grade=override_grade,
        )
        successful_grades += 1
        logging.info(f"Successfully processed and graded: {filename}")

    logging.info("--- Processing Complete ---")
    logging.info(
        f"Total files found: {len(os.listdir(INPUT_FOLDER))}"
    )  # This will count folders too, refine if needed
    logging.info(f"Files attempted for processing: {processed_files}")
    logging.info(f"Successfully graded: {successful_grades}")
    logging.info(f"Reports saved in: {OUTPUT_FOLDER}")
    logging.info(f"Log file saved at: {LOG_FILE}")

    if summary_entries:
        summary_path = os.path.join(OUTPUT_FOLDER, SUMMARY_FILE)
        try:
            with open(summary_path, "w", encoding="utf-8") as sf:
                sf.write("student,total_points,grade\n")
                for ident, points, grade in summary_entries:
                    sf.write(f"{ident},{points},{grade}\n")
            logging.info(f"Summary saved to: {summary_path}")
        except Exception as e:
            logging.error(f"Failed to write summary file: {e}")


if __name__ == "__main__":
    main()
